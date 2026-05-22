"""
tds_cover.py — Post-process a pandoc-generated TDS DOCX to add:
  • Page 1: HPDF logo + document title + doc number + date
  • Page 2: DOE contract attribution + government disclaimer

Called by tds_render.py after pandoc produces the body DOCX.
"""

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── Constants ─────────────────────────────────────────────────────────────────

DOE_ATTRIBUTION = (
    "This material is based upon work supported by the U.S. Department of Energy, "
    "Office of Science, Office of Advanced Scientific Computing Research under contract"
)
DOE_CONTRACT = "DE-AC05-06OR23177"
DOE_DISCLAIMER = (
    "This report was prepared as an account of work sponsored by an agency of the "
    "United States Government. Neither the United States Government nor any agency "
    "thereof, nor any of their employees, makes any warranty, express or implied, or "
    "assumes any legal liability or responsibility for the accuracy, completeness, or "
    "usefulness of any information, apparatus, product, or process disclosed, or "
    "represents that its use would not infringe privately owned rights. Reference "
    "herein to any specific commercial product, process, or service by trade name, "
    "trademark, manufacturer, or otherwise, does not necessarily constitute or imply "
    "its endorsement, recommendation, or favoring by the United States Government or "
    "any agency thereof. The views and opinions of authors expressed herein do not "
    "necessarily state or reflect those of the United States Government or any agency "
    "thereof."
)

# Logo width in the output: ~4.33 inches (matching the source template)
LOGO_WIDTH = Inches(4.33)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _format_doc_number(doc_id: str, slug: str) -> str:
    """
    Format the document number line for the cover page.
    e.g. doc_id='HPDF_TDS_0001', slug='example' → 'HPDF TDS 0001 example'
    Replaces all underscores with spaces in both parts.
    """
    base = doc_id.replace('_', ' ')
    return f"{base} {slug}"


def _format_date(iso_date: str) -> str:
    """
    Parse an ISO date string (YYYY-MM-DD) and return a human-readable form.
    e.g. '2026-05-21' → 'May 21, 2026'
    Falls back to the raw string if parsing fails.
    """
    try:
        dt = datetime.strptime(iso_date.strip(), "%Y-%m-%d")
        return dt.strftime("%B %-d, %Y")
    except (ValueError, TypeError):
        return iso_date


def _make_run_props(para_elem, font_pt=None, bold=False, italic=False,
                    color_hex=None, indent_dxa=None):
    """
    Return an lxml <w:r> element with the given run properties.
    para_elem is only used for namespace context (not modified).
    """
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'),    'Arial')
    fonts.set(qn('w:hAnsi'),    'Arial')
    fonts.set(qn('w:cs'),       'Arial')
    fonts.set(qn('w:eastAsia'), 'Arial')
    rpr.append(fonts)

    if bold:
        b = OxmlElement('w:b')
        rpr.append(b)
        bcs = OxmlElement('w:bCs')
        rpr.append(bcs)

    if italic:
        i = OxmlElement('w:i')
        rpr.append(i)

    if color_hex:
        col = OxmlElement('w:color')
        col.set(qn('w:val'), color_hex)
        rpr.append(col)

    if font_pt:
        sz_val = str(int(font_pt * 2))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), sz_val)
        rpr.append(sz)
        szcs = OxmlElement('w:szCs')
        szcs.set(qn('w:val'), sz_val)
        rpr.append(szcs)

    r.append(rpr)
    return r


def _para_with_text(text, style=None, alignment=None,
                    font_pt=None, bold=False, italic=False,
                    color_hex=None, indent_dxa=None,
                    space_before=None, space_after=None):
    """
    Build and return a python-docx Paragraph element (not yet attached to a document).
    We create it via OxmlElement so we can insert it at arbitrary positions.
    """
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')

    if style:
        ps = OxmlElement('w:pStyle')
        ps.set(qn('w:val'), style)
        ppr.append(ps)

    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        ppr.append(jc)

    if indent_dxa is not None:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(indent_dxa))
        ppr.append(ind)

    if space_before is not None or space_after is not None:
        sp = OxmlElement('w:spacing')
        if space_before is not None:
            sp.set(qn('w:before'), str(space_before))
        if space_after is not None:
            sp.set(qn('w:after'), str(space_after))
        ppr.append(sp)

    p.append(ppr)

    if text:
        r = _make_run_props(p, font_pt=font_pt, bold=bold, italic=italic,
                            color_hex=color_hex)
        t = OxmlElement('w:t')
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    return p


def _page_break_para():
    """Return an lxml paragraph element that forces a page break."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def _disclaimer_table():
    """
    Return an lxml <w:tbl> element containing the DOE disclaimer paragraph.
    5760 DXA wide, centred, full black single-line borders, 9pt text.
    """
    tbl = OxmlElement('w:tbl')

    # Table properties
    tblpr = OxmlElement('w:tblPr')
    tblw = OxmlElement('w:tblW')
    tblw.set(qn('w:w'), '5760')
    tblw.set(qn('w:type'), 'dxa')
    tblpr.append(tblw)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblpr.append(jc)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblpr.append(layout)
    tbl.append(tblpr)

    # Grid
    grid = OxmlElement('w:tblGrid')
    gc = OxmlElement('w:gridCol')
    gc.set(qn('w:w'), '5760')
    grid.append(gc)
    tbl.append(grid)

    # Single row / single cell
    tr = OxmlElement('w:tr')
    tc = OxmlElement('w:tc')

    # Cell properties: full black borders
    tcpr = OxmlElement('w:tcPr')
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '6')
        b.set(qn('w:color'), '000000')
        b.set(qn('w:space'), '0')
        borders.append(b)
    tcpr.append(borders)
    tc.append(tcpr)

    # Cell paragraph with disclaimer text
    p = OxmlElement('w:p')
    ppr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    jc2 = OxmlElement('w:jc')
    jc2.set(qn('w:val'), 'both')
    ppr.append(spacing)
    ppr.append(jc2)
    p.append(ppr)

    r = _make_run_props(p, font_pt=9)
    t = OxmlElement('w:t')
    t.text = DOE_DISCLAIMER
    r.append(t)
    p.append(r)
    tc.append(p)
    tr.append(tc)
    tbl.append(tr)

    return tbl


# ── Public API ────────────────────────────────────────────────────────────────

def prepend_cover(docx_path: Path, doc_id: str, slug: str,
                  last_updated: str, logo_path: Path) -> None:
    """
    Post-process *docx_path* in-place to prepend a cover page and
    DOE disclaimer before the pandoc-generated body.

    The pandoc output starts with a Title-style paragraph (from YAML
    `title:`) followed optionally by a Subtitle paragraph.  This
    function:
      1. Inserts the HPDF logo before the Title paragraph.
      2. Inserts the formatted doc number and date after the last
         initial title-group paragraph (Title / Subtitle).
      3. Inserts a page break to end page 1.
      4. Inserts the DOE attribution + disclaimer table.
      5. Inserts a page break to end page 2, leaving the TOC and
         body content to follow naturally on page 3.

    Args:
        docx_path:    Path to the pandoc-generated DOCX (modified in place).
        doc_id:       e.g. 'HPDF_TDS_0001'
        slug:         e.g. 'example'
        last_updated: ISO date string from YAML frontmatter, e.g. '2026-05-21'
        logo_path:    Path to hpdf-logo.png
    """
    doc = Document(str(docx_path))
    body = doc.element.body

    doc_number = _format_doc_number(doc_id, slug)
    date_str   = _format_date(last_updated)

    # ── Locate the title-group at the start of the pandoc body ───────────────
    # Pandoc places Title / Subtitle paragraphs at indices 0, 1 (or just 0).
    # We find the last index that belongs to the initial title block.

    title_group_end = -1   # index of the last Title/Subtitle para
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            break
        style_elem = child.find('.//' + qn('w:pStyle'))
        style_val  = style_elem.get(qn('w:val'), '') if style_elem is not None else ''
        if style_val in ('Title', 'Subtitle'):
            title_group_end = i
        else:
            break   # first non-title paragraph ends the block

    if title_group_end == -1:
        # No Title paragraph found — insert at position 0 without assumptions
        insert_after = -1
    else:
        insert_after = title_group_end

    # ── Build cover-page elements ─────────────────────────────────────────────

    # Logo paragraph (inline image, left-aligned, Title style for spacing)
    logo_para_docx = doc.add_paragraph()
    logo_para_docx.style = doc.styles['Title']
    if logo_path and Path(logo_path).exists():
        run = logo_para_docx.add_run()
        run.add_picture(str(logo_path), width=LOGO_WIDTH)
    logo_para_elem = logo_para_docx._element
    # Remove the paragraph from the end of the body (add_paragraph appends it)
    body.remove(logo_para_elem)

    # Doc number paragraph: 16pt, left indent 360 DXA
    doc_num_elem = _para_with_text(
        doc_number, font_pt=16, color_hex='808080', indent_dxa=360,
        space_before=40, space_after=160,
    )

    # Date paragraph: 16pt gray
    date_elem = _para_with_text(
        date_str, font_pt=16, color_hex='808080', indent_dxa=360,
        space_before=0, space_after=200,
    )

    # Page break ends cover page 1
    pb1 = _page_break_para()

    # ── Build disclaimer-page elements ────────────────────────────────────────

    # Attribution sentence (bold 11pt)
    attrib_elem = _para_with_text(DOE_ATTRIBUTION, font_pt=11, bold=True,
                                  space_before=0, space_after=0)

    # Contract number (centered bold 11pt)
    contract_elem = _para_with_text(DOE_CONTRACT, font_pt=11, bold=True,
                                    alignment='center', space_before=0, space_after=0)

    # Blank spacer
    blank_elem = _para_with_text('', space_before=0, space_after=0)

    # Government disclaimer table
    disc_tbl = _disclaimer_table()

    # Page break ends disclaimer page 2
    pb2 = _page_break_para()

    # ── Insert everything into the body ──────────────────────────────────────
    #
    # Target layout (indices after all insertions):
    #   0  logo paragraph
    #   1  Title paragraph  (pandoc, from YAML title:)
    #   2  Subtitle paragraph (pandoc, from YAML subtitle:) — if present
    #   …  doc number, date, PAGE BREAK  ← end of cover page 1
    #      attribution, contract, blank, disclaimer table, PAGE BREAK
    #      TOC  (pandoc)
    #      body content  (pandoc)
    #
    # Strategy: insert in reverse order at (insert_after + 1) so each
    # element shifts already-inserted ones forward; then prepend logo at 0.

    insert_pos = insert_after + 1   # right after the title / subtitle block

    # Insert at the same position in "last-desired-first" order so each
    # new element pushes the previous ones one slot forward.
    # Desired final order at insert_pos: doc_num_elem, date_elem, pb1,
    #   attrib_elem, contract_elem, blank_elem, disc_tbl, pb2
    # → insert them in REVERSE of that desired order.
    for elem in [pb2, disc_tbl, blank_elem, contract_elem, attrib_elem, pb1,
                 date_elem, doc_num_elem]:
        body.insert(insert_pos, elem)

    # Logo always goes at position 0 — pandoc's Title paragraph is at 0
    # (or the first content paragraph if there was no title group).
    body.insert(0, logo_para_elem)

    # ── Insert page break after the TOC ──────────────────────────────────────
    # The TOC is a w:sdt element.  Find it and insert a page-break paragraph
    # immediately after so the body content starts cleanly on a new page.
    children = list(body)
    for i, child in enumerate(children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sdt':
            body.insert(i + 1, _page_break_para())
            break   # only the first sdt — that's the TOC

    # ── Save in place ─────────────────────────────────────────────────────────
    doc.save(str(docx_path))
